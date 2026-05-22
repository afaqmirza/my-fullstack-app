"""
Safe DOCX text/table extraction without loading broken word/media/* references.
Fixes: "There is no item named 'word/media/image1.png' in the archive"
"""
import zipfile
from xml.etree import ElementTree as ET

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"


def _read_document_xml(path: str) -> ET.Element:
    with zipfile.ZipFile(path, "r") as zf:
        if "word/document.xml" not in zf.namelist():
            raise ValueError("Invalid DOCX file: missing document body.")
        return ET.fromstring(zf.read("word/document.xml"))


def _paragraph_text(element) -> str:
    parts = []
    for node in element.iter(f"{W}t"):
        if node.text:
            parts.append(node.text)
        if node.tail:
            parts.append(node.tail)
    return "".join(parts).strip()


def extract_docx_text_safe(path: str) -> str:
    """Extract all paragraph text from a DOCX without touching embedded images."""
    try:
        root = _read_document_xml(path)
        lines = []
        for para in root.iter(f"{W}p"):
            line = _paragraph_text(para)
            if line:
                lines.append(line)
        if lines:
            return "\n".join(lines)
    except zipfile.BadZipFile as e:
        raise ValueError(f"Invalid or corrupted DOCX file: {e}") from e
    except Exception:
        pass

    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(path)
        lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        if lines:
            return "\n".join(lines)
    except Exception as e:
        raise ValueError(
            "Could not read this Word file. It may have broken embedded images. "
            "Open it in Microsoft Word, use Save As → new .docx, or upload PDF/TXT instead. "
            f"({e})"
        ) from e

    return ""


def extract_docx_tables_safe(path: str) -> list:
    """Extract tables as list of rows (list of cell strings)."""
    tables_data = []
    try:
        root = _read_document_xml(path)
        for tbl in root.iter(f"{W}tbl"):
            table_data = []
            for tr in tbl.findall(f"{W}tr"):
                row_data = []
                for tc in tr.findall(f"{W}tc"):
                    cell_parts = []
                    for t in tc.iter(f"{W}t"):
                        if t.text:
                            cell_parts.append(t.text)
                    row_data.append("".join(cell_parts).strip())
                table_data.append(row_data)
            if table_data:
                tables_data.append(table_data)
        if tables_data:
            return tables_data
    except Exception:
        pass

    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(path)
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text.strip() for cell in row.cells])
            if table_data:
                tables_data.append(table_data)
    except Exception:
        pass

    return tables_data
